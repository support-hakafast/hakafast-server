# -*- coding: utf-8 -*-
"""Generate HAKAFAST Hebrew guide as DOCX."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import os
ROOT = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(ROOT, "docx", "HAKAFAST-Masslul-Guide.he.docx")


def set_rtl(paragraph, align_right=True):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_rtl(p)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1) if level == 1 else RGBColor(0x15, 0x65, 0xC0)
    return p


def add_para(doc, text, bold=False, size=12):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        set_rtl(p)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        set_rtl(p)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            set_rtl(p)
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                set_rtl(p)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)
    doc.add_paragraph()


def add_note(doc, text, kind="note"):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True
    return p


doc = Document()
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21)
section.left_margin = Cm(2)
section.right_margin = Cm(2)

# Cover
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p, align_right=False)
r = p.add_run("HAKAFAST — מדריך מסלול")
r.bold = True
r.font.size = Pt(26)
r.font.name = "Arial"
r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p2, align_right=False)
r2 = p2.add_run("תור נהגים · קבלה (Reception) · Rentix · התקנה ותפעול")
r2.font.size = Pt(14)
r2.font.name = "Arial"

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p3, align_right=False)
r3 = p3.add_run("גרסת מסמך: יוני 2026")
r3.font.size = Pt(11)
r3.font.name = "Arial"
r3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

add_heading(doc, "1. מה זה HAKAFAST?")
add_para(doc, "HAKAFAST הוא מערכת ניהול מקצים וזמנים חיים למסלולי קארטינג. היא מנהלת את מה שקורה במסלול: תור נהגים, שיבוץ לקארטים, יציאה מפיטס, הקפות, תוצאות, מעבר בין מקצים — ולא את מערכת הקופה או האתר.")
add_bullets(doc, [
    "פאנל ניהול (Admin) — שליטה מלאה: תור, שיבוץ, פיטס, הגדרות מקצה, ייצוא תוצאות.",
    "זמנים חיים (Live Timing) — תצוגה למסך גדול / QR / טלוויזיה במסלול.",
    "קבלה (Reception) — מסך פשוט לרישום נהגים לתור (טאבלט בדלפק).",
])

add_heading(doc, "2. ארכיטקטורה — מי עושה מה", level=2)
add_heading(doc, "2.1 מנוע המקצים", level=3)
add_para(doc, "כל המידע החי — תור, מקצה נוכחי, קארטים בפיטס, הקפות — נשמר ב-workspace מקומי (קובץ JSON תחת %ProgramData%\\HAKAFAST\\workspaces\\). זה המנוע שמנהל את היום-יום במסלול.")
add_heading(doc, "2.2 TranX / טרנספונדרים", level=3)
add_para(doc, "חומרת הזמנים (AMB TranX, לולאות מגנטיות, טרנספונדרים) שולחת אירועים ישירות ל-HAKAFAST. Rentix לא מנהל זמנים — TranX → HAKAFAST.")
add_heading(doc, "2.3 PostgreSQL", level=3)
add_para(doc, "בסיס הנתונים בענן משמש בעיקר לנתיב legacy ולגיבוי פרופיל מסלול. הוא אינו מנוע המקצים החי.")
add_heading(doc, "2.4 Rentix (אופציונלי)", level=3)
add_para(doc, "Rentix הוא מערכת CMS חיצונית להזמנות ותשלומים באינטרנט. HAKAFAST מתחבר אליה רק אם המסלול כבר מוכר sessions דרך Rentix.")

add_heading(doc, "3. תור נהגים — המושג המרכזי")
add_para(doc, "תור נהגים (driverQueue) הוא רשימה של נהגים שממתינים למקצה הבא. כל נהג: שם, טלפון (אופציונלי), רמה (Amateur/Master/Pro), שם קבוצה (בסיבולת), מקור (ידני/reception/rentix).")
add_para(doc, "התור הוא אותו תור בכל המערכת. מנהל המסלול לוחץ \"שבץ קארטים\" והמערכת לוקחת נהגים מהתור ומשבצת לקארטים פנויים.")

add_heading(doc, "4. קבלה (Reception) — מה זה ולמי")
add_heading(doc, "4.1 מה זה?", level=3)
add_para(doc, "מסך Reception הוא ממשק קל לרישום נהגים שמגיעים למסלול — טאבלט בדלפק הקבלה. כתובת: http://127.0.0.1:5000/reception/{track-slug}")
add_heading(doc, "4.2 מה אפשר לעשות?", level=3)
add_bullets(doc, [
    "לראות כמה נהגים בתור ומספר המקצה הנוכחי",
    "להוסיף נהג: שם, טלפון, רמה",
    "להסיר נהג מהתור",
])
add_heading(doc, "4.3 מה אי אפשר?", level=3)
add_bullets(doc, [
    "לשבץ קארטים, להתחיל/לעצור מקצה, לנהל פיטס",
    "לקבל תשלום, לנהל הזמנות מהאינטרנט, זמנים חיים מפורטים",
])
add_note(doc, "בקיצור: Reception = \"רישום שם לתור\". כל השאר קורה ב-Admin.")

add_table(doc, ["פעולה", "Endpoint", "תיאור"], [
    ["קריאת מצב", "GET /api/reception/state", "תור, מספר מקצה, האם מקצה פעיל"],
    ["הוספת נהג", "POST /api/reception/drivers", "name, phone, driver_level"],
    ["הסרת נהג", "DELETE /api/reception/drivers/:index", "לפי מיקום בתור"],
])

add_heading(doc, "5. תור נהגים בפאנל הניהול (Admin)")
add_bullets(doc, [
    "הוספה / עריכה / הסרה של נהגים",
    "שיבוץ אוטומטי — \"שבץ קארטים\"",
    "ניהול פיטס, יציאה, הקפות, סיום מקצה, מעבר למקצה הבא",
    "ייצוא CSV/PDF בסיום",
])
add_para(doc, "Reception ו-Admin משתפים את אותו מערך driverQueue בזמן אמת.")

add_heading(doc, "6. Rentix — מה זה ולמה בכלל צריך")
add_heading(doc, "6.1 מה זה Rentix?", level=3)
add_para(doc, "Rentix (cms.rentix.biz) הוא מערכת חיצונית: אתר הזמנות, תשלום מראש, ניהול מחירים, CRM.")
add_heading(doc, "6.2 מה HAKAFAST עושה עם Rentix?", level=3)
add_para(doc, "HAKAFAST לא מחליף את Rentix. הוא רק מקבל שמות נהגים ששילמו ומכניס אותם ל-driverQueue — כמו שהקבלה מוסיפה שם ידנית.")
add_heading(doc, "6.3 איך הנתונים נכנסים?", level=3)
add_bullets(doc, [
    "Webhook: Rentix → POST /api/webhooks/rentix → הוספה לתור",
    "Sync: POST /api/webhooks/rentix/sync → משיכת הזמנות ששולמו (24 שעות)",
])
add_heading(doc, "6.4 תוצאות חזרה (אופציונלי)", level=3)
add_para(doc, "בסיום מקצה — שליחת תוצאות ל-resultsWebhookUrl. שימושי להצגה באתר, לא חובה לתפעול.")
add_note(doc, "חשוב: Rentix = הזמנות ותשלומים. HAKAFAST = מקצים וזמנים. Rentix רק מזין שמות לתור אם כבר משתמשים בו.")

add_heading(doc, "7. השוואה: Reception / Admin / Rentix")
add_table(doc, ["", "Reception", "Admin", "Rentix"], [
    ["מטרה", "רישום walk-in לתור", "ניהול מלא של המקצה", "הזמנות ותשלום באינטרנט"],
    ["מי משתמש", "פקיד קבלה", "מנהל מקצה", "לקוח + משרד"],
    ["תשלום", "לא", "לא", "כן"],
    ["שיבוץ קארטים", "לא", "כן", "לא"],
    ["זמנים חיים", "לא", "כן", "לא"],
    ["חובה?", "מומלץ לדלפק", "חובה", "רק אם מוכרים דרך Rentix"],
    ["אותו תור?", "כן", "כן", "כן — נכנס ל-driverQueue"],
])

add_heading(doc, "8. זרימת עבודה יומית")
add_heading(doc, "8.1 בלי Rentix", level=3)
add_numbered(doc, [
    "לקוח משלם בקופה (POS / מזומן)",
    "פקיד קבלה מוסיף שם ב-Reception",
    "מנהל מקצה שובץ קארטים ב-Admin",
    "נהגים יוצאים (TranX) → Live Timing",
    "סיום מקצה → מקצה הבא",
])
add_heading(doc, "8.2 עם Rentix", level=3)
add_numbered(doc, [
    "לקוח הזמין ושילם באתר → webhook → תור",
    "Walk-in עדיין נרשם ב-Reception",
    "שיבוץ ומקצה כרגיל",
    "אופציונלי: תוצאות חוזרות ל-Rentix",
])

add_heading(doc, "9. מתי כן צריך Rentix ומתי לא")
add_heading(doc, "לא צריך Rentix אם:", level=3)
add_bullets(doc, [
    "כל התשלומים בקופה / POS",
    "אין אתר Rentix",
    "רוצים רק HAKAFAST + Reception + TranX",
])
add_heading(doc, "כן שווה לחבר Rentix אם:", level=3)
add_bullets(doc, [
    "כבר מוכרים דרך Rentix ורוצים שמות אוטומטיים לתור",
    "רוצים להימנע מהקלדה כפולה",
    "רוצים תוצאות באתר הלקוח",
])
add_note(doc, "המלצה: התחילו ב-Reception + Admin. הוסיפו Rentix רק כשיש הזמנות אונליין פעילות.")

add_heading(doc, "10. התקנה במסלול (ישראל)")
add_numbered(doc, [
    "מחשב Windows (KIOSK) + TranX 160 + טרנספונדרים + לולאות",
    "Setup Wizard — שם מסלול, workspace",
    "Track Planner — שעות, מחיר, משך session",
    "Admin / Reception / Live Timing על מסכים נפרדים",
])
add_para(doc, "קבצים: %ProgramData%\\HAKAFAST\\install.json, workspaces\\*.json, ambTranx160.js")

add_heading(doc, "11. TranX וטרנספונדרים")
add_table(doc, ["אירוע", "Endpoint", "מתי"], [
    ["יציאה מפיטס", "POST /api/transponder/pit-exit", "חציית pit-out"],
    ["כניסה לפיטס", "POST /api/transponder/pit-entry", "חזרה לפיטס"],
    ["הקפה", "POST /api/transponder/lap", "קו סיום"],
])

add_heading(doc, "12. צ'קליסט יום פתיחה")
add_bullets(doc, [
    "HAKAFAST service רץ",
    "TranX מחובר (npm run verify:amb)",
    "Admin + Reception + Live Timing פתוחים",
    "תור מוכן, קארטים במצב תקין",
    "(אופציונלי) Rentix status: GET /api/webhooks/rentix/status",
])

add_heading(doc, "13. שאלות נפוצות")
add_para(doc, "למה גם Reception וגם תור ב-Admin? — אותו תור, שני ממשקים: Reception לפקיד, Admin למנהל מקצה.", bold=False)
add_para(doc, "האם הזמנות Rentix = תור HAKAFAST? — לא. הזמנה = מכירה. תור = מי ממתין למקצה. Rentix רק מעתיק שם.")
add_para(doc, "מה בלי Rentix? — הכל עובד. Reception + Admin + TranX = מסלול מלא.")
add_para(doc, "Reception מחליף קופה? — לא. תשלום נשאר בקופה / POS / Rentix.")

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")
