# -*- coding: utf-8 -*-
"""Generate HAKAFAST business one-pager as DOCX."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

import os
ROOT = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(ROOT, "docx", "HAKAFAST-Business-OnePager.he.docx")


def set_rtl(paragraph, align_right=True):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_rtl(p)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    return p


def para(doc, text, bold=False, size=10.5):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def bullets(doc, items, size=10.5):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        set_rtl(p)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)


def numbered(doc, items, size=10.5):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        set_rtl(p)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)


def table(doc, headers, rows, size=9.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            set_rtl(p)
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(size)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                set_rtl(p)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(size)


doc = Document()
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21)
sec.left_margin = Cm(1.5)
sec.right_margin = Cm(1.5)
sec.top_margin = Cm(1.2)
sec.bottom_margin = Cm(1.2)

p = doc.add_paragraph()
set_rtl(p)
r = p.add_run("HAKAFAST — One-Pager עסקי")
r.bold = True
r.font.size = Pt(20)
r.font.name = "Arial"
r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

para(doc, "מודל תמחור לשוק ישראלי · B2B · On-Premise · יוני 2026", size=10)

para(doc, "המלצה מרכזית: מכירת רישיון לכל החיים למסלול (Per Site) — כמו Rentix — לא מנוי חודשי, לא אחוז מעסקה. שוק קטן (6–8 מסלולים) = מכירה ישירה, התקנה, יחסים ארוכי טווח.", bold=True)

heading(doc, "מה HAKAFAST מוכר / לא מוכר", 2)
table(doc, ["מוכר", "לא מוכר"], [
    ["תזמון חי, פיטס, שיבוץ, Live, Reception, TranX, 3 שפות, offline", "קופה, POS, הזמנות אונליין*, CRM"],
])
para(doc, "* חיבור Rentix — תוסף אופציונלי.", size=9)

heading(doc, "מיקום בשוק — ליד Rentix", 2)
table(doc, ["", "Rentix", "HAKAFAST"], [
    ["תפקיד", "קופה, CRM, לפעמים אתר", "מקצים, פיטס, זמנים"],
    ["מודל", "רישיון לכל החיים", "רישיון לכל החיים (מומלץ)"],
    ["TranX", "לא מנהל timing", "ליבת המוצר"],
])

heading(doc, "מודל תמחור — 3 רכיבים", 2)
table(doc, ["רכיב", "סוג", "תיאור"], [
    ["1. רישיון לכל החיים", "חד-פעמי", "Site אחד: Admin + Live + Reception + TranX"],
    ["2. התקנה והדרכה", "חד-פעמי", "Setup, חומרה, יום הדרכה"],
    ["3. תמיכה ועדכונים", "שנתי (אופצ.)", "~15%–20% מערך הרישיון"],
])

heading(doc, "חבילות רישיון", 2)
table(doc, ["חבילה", "קארטים", "כלול"], [
    ["Standard", "עד 15", "Admin, Live, Reception, TranX, CSV/PDF"],
    ["Pro", "16–25", "+ סיבולת, ספרינט, תכנון יום, 2 מסכי Live"],
    ["Enterprise", "26+", "+ endurance rules, SLA, רב-מתחם"],
])
para(doc, "מחיר סופי בהצעה לפי קaarטים, מסכים, מרחק התקנה. ללא מחירון ציבורי.", size=9)

heading(doc, "תוספים (חד-פעמי)", 2)
bullets(doc, [
    "מסך Live / Reception נוסף",
    "חיבור Rentix",
    "ביקור התקנה · שדרוג חבילה",
])

heading(doc, "למה לא SaaS / לא % מעסקה", 2)
bullets(doc, [
    "שוק קטן — 6–8 מסלולים",
    "ציפייה מקומית — רישיון לכל החיים כמו Rentix",
    "מוצר offline — ערך במחשב במסלול",
    "5–6 לקוחות × (רישיון + התקנה) = מודל ישיר",
])

heading(doc, "זרימת מכירה", 2)
numbered(doc, [
    "דמו (kart-demo)",
    "פגישה במסלול",
    "הצעה: רישיון + התקנה + תמיכה (אופצ.)",
    "חוזה · התקנה · הדרכה",
])

para(doc, "support.hakafast@gmail.com · מחירים סופיים בהצעה חתומה", size=9)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"Saved: {OUT}")
