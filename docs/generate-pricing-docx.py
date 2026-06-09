# -*- coding: utf-8 -*-
"""Generate HAKAFAST pricing tiers DOCX with concrete ILS prices."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(ROOT, "docx", "HAKAFAST-Pricing-Tiers.he.docx")

# ── מחירון (₪, לפני מע"מ) ──────────────────────────────────────
LICENSE = {
    "Standard": {"karts": "עד 15", "price": 28_000},
    "Pro": {"karts": "16–25", "price": 42_000},
    "Enterprise": {"karts": "26+", "price": 58_000},
}
SERVICES = {
    "install": 8_500,
    "training_day": 4_500,
}
SUPPORT_PCT = 0.18  # שנתי, אופציונלי
ADDONS = {
    "live_screen": 3_500,
    "reception": 2_000,
    "rentix": 4_500,
    "travel_outside": 2_500,
}
UPGRADE = {
    "Standard→Pro": 14_000,
    "Pro→Enterprise": 16_000,
    "Standard→Enterprise": 30_000,
}
VAT = 0.18


def ils(n):
    return f"₪{n:,}".replace(",", ",")


def set_rtl(paragraph, align_right=True):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def title(doc, text, size=20, color=0x0D47A1):
    p = doc.add_paragraph()
    set_rtl(p)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
    return p


def heading(doc, text):
    p = doc.add_heading(text, level=2)
    set_rtl(p)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    return p


def para(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    set_rtl(p)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    return p


def table(doc, headers, rows, size=10.5):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            set_rtl(p)
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(size)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                set_rtl(p)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(size)


def example_total(doc, pkg_name):
    lic = LICENSE[pkg_name]["price"]
    sub = lic + SERVICES["install"] + SERVICES["training_day"]
    support = int(round(lic * SUPPORT_PCT))
    vat = int(round(sub * VAT))
    total = sub + vat
    heading(doc, f"דוגמה — חבילת {pkg_name} (התקנה מלאה, ללא תמיכה שנתית)")
    table(doc, ["פריט", "מחיר"], [
        [f"רישיון {pkg_name} לכל החיים", ils(lic)],
        ["התקנה + TranX", ils(SERVICES["install"])],
        ["הדרכה (יום)", ils(SERVICES["training_day"])],
        ['סה"כ לפני מע"מ', ils(sub)],
        ['מע"מ 18%', ils(vat)],
        ['סה"כ לתשלום', ils(total)],
    ])
    para(doc, f"עם תמיכה שנתית (+{ils(support)} לפני מע\"מ): {ils(total + int(round(support * (1 + VAT))))} כולל מע\"מ", size=10)


os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc = Document()
sec = doc.sections[0]
sec.left_margin = Cm(1.8)
sec.right_margin = Cm(1.8)

title(doc, "HAKAFAST — מחירון רישיון")
para(doc, "רישיון לכל החיים · מסלול אחד (Site) · מחירים בש\"ח לפני מע\"מ · יוני 2026")
para(doc, "support.hakafast@gmail.com", size=10)

heading(doc, "1. רישיון לכל החיים — לפי חבילה")
table(doc, ["חבילה", "קaarטים", "מחיר", "כלול"], [
    ["Standard", LICENSE["Standard"]["karts"], ils(LICENSE["Standard"]["price"]),
     "Admin · Live · Reception · TranX · CSV/PDF · 1 מסך Live"],
    ["Pro", LICENSE["Pro"]["karts"], ils(LICENSE["Pro"]["price"]),
     "Standard + סיבולת · ספרינט · תכנון יום · 2 מסכי Live"],
    ["Enterprise", LICENSE["Enterprise"]["karts"], ils(LICENSE["Enterprise"]["price"]),
     "Pro + endurance rules · SLA · תמיכה מועדפת"],
])

heading(doc, "2. שירותים חד-פעמיים")
table(doc, ["שירות", "מחיר"], [
    ["התקנה, הגדרה וחיבור TranX 160", ils(SERVICES["install"])],
    ["הדרכת צוות — יום אחד (מנהל מקצה + קבלה)", ils(SERVICES["training_day"])],
])

heading(doc, "3. תמיכה ועדכונים — שנתי (אופציונלי)")
table(doc, ["חבילה", "מחיר שנתי (18% מערך הרישיון)"], [
    ["Standard", ils(int(LICENSE["Standard"]["price"] * SUPPORT_PCT))],
    ["Pro", ils(int(LICENSE["Pro"]["price"] * SUPPORT_PCT))],
    ["Enterprise", ils(int(LICENSE["Enterprise"]["price"] * SUPPORT_PCT))],
])
para(doc, "כולל: תיקונים, גרסאות minor, תמיכה טלפון / מייל / וואטסאפ בשעות עסקים.", size=10)

heading(doc, "4. תוספים (חד-פעמי)")
table(doc, ["תוסף", "מחיר"], [
    ["מסך Live Timing נוסף", ils(ADDONS["live_screen"])],
    ["מסך Reception נוסף", ils(ADDONS["reception"])],
    ["חיבור Rentix / webhook", ils(ADDONS["rentix"])],
    ["נסיעה מחוץ למרכז (גוש דן / חיפה) — תוספת", ils(ADDONS["travel_outside"])],
])

heading(doc, "5. שדרוג חבילה")
table(doc, ["שדרוג", "תוספת (הפרש רישיון)"], [
    [k, ils(v)] for k, v in UPGRADE.items()
])

heading(doc, "6. תנאי תשלום")
para(doc, "50% בחתימת הזמנה · 50% ב-Go-Live (התקנה + הדרכה הושלמו) · חשבונית מס · מע\"מ 18% על כל סכום.")

example_total(doc, "Pro")

para(doc, "מחירון פנימי / מכירות. הצעה סופית למסלול ספציפי יכולה לכלול הנחת Early Adopter — לפי שיקול דעת.", size=9)

doc.save(OUT)
print(f"Saved: {OUT}")
