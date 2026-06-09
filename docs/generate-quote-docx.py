# -*- coding: utf-8 -*-
"""Generate HAKAFAST quote template as DOCX."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(ROOT, "docx", "HAKAFAST-Quote-Template.he.docx")


def set_rtl(paragraph, align_right=True):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    if align_right:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def heading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    set_rtl(p)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    return p


def para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    return p


def bullets(doc, items, size=10.5):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        set_rtl(p)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)


def numbered(doc, items, size=10.5):
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        set_rtl(p)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)


def table(doc, headers, rows, size=10):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            set_rtl(p)
            for run in p.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(size)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
            for p in t.rows[ri + 1].cells[ci].paragraphs:
                set_rtl(p)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(size)


os.makedirs(os.path.dirname(OUT), exist_ok=True)

doc = Document()
sec = doc.sections[0]
sec.page_height = Cm(29.7)
sec.page_width = Cm(21)
sec.left_margin = Cm(2)
sec.right_margin = Cm(2)

p = doc.add_paragraph()
set_rtl(p)
r = p.add_run("HAKAFAST — הצעת מחיר")
r.bold = True
r.font.size = Pt(22)
r.font.name = "Arial"
r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

para(doc, "רישיון תוכנה לכל החיים · מסלול קארטינג", size=11)
para(doc, "support.hakafast@gmail.com", size=10)

para(doc, "תאריך: ___ / ___ / 2026          מס' הצעה: HF-2026-___")
para(doc, "לכבוד: _________________________ (שם המסלול / חברה)")
para(doc, "איש קשר: _________________________ (שם, טלפון, מייל)")

heading(doc, "1. מטרת ההצעה")
para(doc, "רכישת רישיון HAKAFAST לכל החיים (Per Site) למסלול אחד — התקנה, TranX 160, הדרכה. תפעול מקומי ללא תלות באינטרנט.")

heading(doc, "2. היקף הרישיון")
bullets(doc, [
    "מיקום (Site) אחד: _________________________",
    "כמות קaarטים: ___",
    "חבילה: Standard / Pro / Enterprise",
])

heading(doc, "3. מה כלול ברישיון")
table(doc, ["רכיב", "כלול"], [
    ["Admin", "פיטס, שיבוץ, מקצים, CSV/PDF"],
    ["Live Timing", "1 מסך (ניתן להוסיף)"],
    ["Reception", "רישום נהגים לתור"],
    ["TranX 160", "יציאה, הקפות, פיטס"],
    ["שפות", "עברית, ערבית, אנגלית"],
    ["עדכונים", "12 חודשים ראשונים — minor"],
])
para(doc, "לא כלול: קופה, POS, הזמנות אונליין.", size=10)

heading(doc, "4. פירוט מחיר")
table(doc, ["#", "פריט", "פירוט", "מחיר (₪)"], [
    ["1", "רישיון לכל החיים", "חבילה Pro · עד 25 קaarטים", "________"],
    ["2", "התקנה + TranX", "Setup, workspace, בדיקות", "________"],
    ["3", "הדרכה (יום)", "מנהל מקצה + קבלה", "________"],
    ["4", "תמיכה שנה 1 (אופצ.)", "טלפון / מייל / עדכונים", "________"],
    ["", "סה\"כ לפני מע\"מ", "", "________"],
    ["", "מע\"מ 18%", "", "________"],
    ["", "סה\"כ כולל מע\"מ", "", "________"],
])

heading(doc, "5. תוספים אופציונליים")
table(doc, ["תוסף", "מחיר (₪)"], [
    ["מסך Live נוסף", "________"],
    ["Reception נוסף", "________"],
    ["חיבור Rentix", "________"],
    ["נסיעות מחוץ לאזור", "לפי actual"],
])

heading(doc, "6. תנאי תשלום")
bullets(doc, [
    "50% עם חתימת הזמנה",
    "50% עם Go-Live (התקנה + הדרכה)",
    "חשבונית מס לכל תשלום",
])

heading(doc, "7. לוח זמנים")
table(doc, ["שלב", "משך"], [
    ["התקנה + הגדרות", "1–2 ימים"],
    ["TranX + בדיקות", "1 יום"],
    ["הדרכה + Go-Live", "1 יום"],
])

heading(doc, "8. תנאי רישיון (תמצית)")
bullets(doc, [
    "רישיון לכל החיים — Site אחד, ללא העברה ללא אישור",
    "קוד המקור בבעלות HAKAFAST — זכות שימוש perpetual למסלול",
    "גיבוי — באחריות המסלול",
    "תמיכה משנה 2 — חידוש אופציונלי",
])

heading(doc, "9. תוקף")
para(doc, "הצעה תקפה 30 יום מתאריך ההצעה.")

heading(doc, "10. אישור")
para(doc, "חתימת המסלול: _________________   תאריך: _________")
para(doc, "HAKAFAST: _________________   תאריך: _________")
para(doc, "שליחה ל: support.hakafast@gmail.com", size=10)

doc.save(OUT)
print(f"Saved: {OUT}")
