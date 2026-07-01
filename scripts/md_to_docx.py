#!/usr/bin/env python3
"""Convert MADRICH-HAKAFAT-ISRAEL.md to a formatted DOCX (Hebrew RTL)."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_MD = ROOT / "installer" / "MADRICH-HAKAFAT-ISRAEL.md"
DEFAULT_OUT = ROOT / "installer" / "MADRICH-HAKAFAT-ISRAEL.docx"

NAVY = RGBColor(0, 0, 128)
TEAL = RGBColor(0, 128, 128)
MUTED = RGBColor(100, 116, 139)


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    p_pr.append(bidi)


def set_run_font(run, size=11, bold=False, color=None, name="Arial"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:cs"), name)
    r_pr.insert(0, r_fonts)


def add_para(doc, text, size=11, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        set_rtl(p, WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(line)
        set_run_font(run, size=9, name="Consolas")
        run.font.color.rgb = RGBColor(30, 41, 59)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "F1F5F9")
        p._element.get_or_add_pPr().append(shading)


def parse_table(lines):
    rows = []
    for line in lines:
        if not line.strip().startswith("|"):
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.match(r"^[-:\s]+$", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = text
            for p in cell.paragraphs:
                set_rtl(p)
                for run in p.runs:
                    set_run_font(run, size=10, bold=(ri == 0))
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # Title page
    add_para(doc, "HAKAFAST", size=28, bold=True, color=NAVY, space_after=12)
    add_para(doc, "מדריך התקנה במסלול בישראל", size=20, bold=True, color=TEAL, space_after=8)
    add_para(doc, "מסמך עבודה בשלבים — צוות התקנה, מנהל מסלול ותמיכה", size=12, color=MUTED, space_after=4)
    add_para(doc, "גרסה 1.0 · Windows 10/11 (64-bit)", size=11, color=MUTED, space_after=24)
    add_para(doc, "support.hakafast@gmail.com · https://hakafast.com", size=10, color=MUTED, space_after=0)
    doc.add_page_break()

    # TOC placeholder
    add_para(doc, "תוכן עניינים", size=16, bold=True, color=NAVY, space_after=12)
    toc_items = re.findall(r"^\d+\.\s+\[([^\]]+)\]", text, re.MULTILINE)
    for i, item in enumerate(toc_items, 1):
        add_para(doc, f"{i}. {item}", size=11, space_after=4)
    doc.add_page_break()

    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.startswith("# "):
            add_para(doc, line[2:].strip(), size=22, bold=True, color=NAVY, space_after=10)
        elif line.startswith("## "):
            add_para(doc, line[3:].strip(), size=16, bold=True, color=TEAL, space_after=8)
        elif line.startswith("### "):
            add_para(doc, line[4:].strip(), size=13, bold=True, color=NAVY, space_after=6)
        elif line.strip().startswith("|") and "|" in line:
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_table(doc, parse_table(table_lines))
            continue
        elif line.strip().startswith("- [ ]"):
            add_para(doc, "☐ " + line.strip()[5:].strip(), size=11, space_after=4)
        elif line.strip().startswith("- "):
            add_para(doc, "• " + line.strip()[2:].strip(), size=11, space_after=4)
        elif re.match(r"^\d+\.\s", line.strip()):
            add_para(doc, line.strip(), size=11, space_after=4)
        elif line.strip().startswith(">"):
            add_para(doc, line.strip().lstrip(">").strip(), size=10, color=MUTED, space_after=6)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.strip():
            clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line.strip())
            clean = clean.replace("`", "")
            add_para(doc, clean, size=11, space_after=6)

        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Created: {out_path}")


if __name__ == "__main__":
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_MD
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else DEFAULT_OUT
    convert(md, out)
